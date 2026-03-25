import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO

# --- 初始化防呆與狀態 (Session State) ---
if 'setup_complete' not in st.session_state:
    st.session_state.setup_complete = False
if 'df' not in st.session_state:
    st.session_state.df = None
if 'date_str' not in st.session_state:
    st.session_state.date_str = None
if 'location' not in st.session_state:
    st.session_state.location = ""
if 'attendees' not in st.session_state:
    st.session_state.attendees = [] 
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False 
if 'excel_data' not in st.session_state:
    st.session_state.excel_data = None
if 'word_data' not in st.session_state:
    st.session_state.word_data = None
if 'final_df' not in st.session_state:
    st.session_state.final_df = None

# --- 資料庫設定 ---
TEAM_MEMBERS = [
    "丁秋吟", "王永慶", "王銓德", "王志文", "王家業", "朱家樺", "江旼珀", "吳上苑", "吳宜汶", "呂恩昕", 
    "呂淑惠", "李國誥", "李榮斌", "李穎裕", "阮智偉", "周昆佑", "周志暐", "周志祥", "林志嶸", "林永松", 
    "林志佳", "劉柔君", "林佳宏", "林國芳", "林婉茹", "林智偉", "林華盛", "林瑞華", "林瑞燿", "林嘉信", 
    "邱保銘", "邱信培", "徐琮凱", "柯富強", "洪偉倫", "胡中興", "胡耀仁", "夏進通", "涂旻聖", "張進源", 
    "張文男", "張世明", "張仕欣", "張正中", "張百江", "張孟哲", "張錦升", "張勝富", "張路雄", "張鈞銘", 
    "張仕宗", "張聰捷", "梁善鈞", "郭政富", "陳文政", "陳仕明", "陳怡如", "陳瑜玲", "陳冠良", "陳春文", 
    "陳敏訓", "陳盛宏", "陳進忠", "陳佳忠", "曾雅惠", "游正豪", "游育民", "游振和", "黃世政", "黃明興", 
    "黃信隆", "黃冠霖", "黃堃珉", "黃賀進", "黃麗萍", "楊俊逸", "魏捷祥", "楊閔森", "廖崇凱", "廖竣傑", 
    "廖致維", "熊致堯", "蔡芷纭", "蔡玉雯", "蔡榮峰", "鄭銘宗", "蕭大勛", "蕭鈺潔", "謝仁政", "魏志龍", 
    "藍慧真", "詹獻睿", "詹孟杰", "曾舜麟", "謝明佑", "羅靜宜", "姜小平", "莊麗雪", "吳慧芳", "蔡瑞賓", 
    "張上觀", "陳素貞"
]

CONSULTANTS = [
    "吳建興", "呂昇印", "孟繁光", "林坤茂", "邱榮家", "徐偉欽", "鄭淵太", "張志州", 
    "張富山", "陳珀升", "陳溪宗", "陳瑋楊", "曾明雄", "詹憲國", "廖宏輝", "廖翊均", "劉邦杰", "劉明煌", 
    "蔡榮祥", "蔡榮華", "周智勤", "謝志忠", "涂欽耀", "張志仲", "蕭森巍", "張銀恭", "王正錄", "曾建勳", 
    "黃智煒", "劉海森", "賴南君", "劉權漢", "游伊君", "陳勇志", "陳文田", "林秋雄", "賴永昌", "劉煉騰", 
    "林錦志", "林明忠", "張哲誠", "詹昆學", "陳彥宏", "許馨云", "張橋語"
]

MEMBER_ROLES = {name: "隊員" for name in TEAM_MEMBERS}
MEMBER_ROLES.update({name: "顧問" for name in CONSULTANTS})

RAW_MEMBERS = [
    "丁秋吟(ㄉㄑㄧ)", "王永慶(ㄨㄩㄑ)", "王銓德(ㄨㄑㄉ)", "王志文(ㄨㄓㄨ)", "王家業(ㄨㄐㄧ)", "朱家樺(ㄓㄐㄏ)", "江旼珀(ㄐㄇㄆ)", "吳上苑(ㄨㄕㄩ)", "吳宜汶(ㄨㄧㄨ)", "呂恩昕(ㄌㄣㄒ)", 
    "呂淑惠(ㄌㄕㄏ)", "李國誥(ㄌㄍㄍ)", "李榮斌(ㄌㄖㄅ)", "李穎裕(ㄌㄧㄩ)", "阮智偉(ㄖㄓㄨ)", "周昆佑(ㄓㄎㄧ)", "周志暐(ㄓㄓㄨ)", "周志祥(ㄓㄓㄒ)", "林志嶸(ㄌㄓㄖ)", "林永松(ㄌㄩㄙ)", 
    "林志佳(ㄌㄓㄐ)", "劉柔君(ㄌㄖㄐ)", "林佳宏(ㄌㄐㄏ)", "林國芳(ㄌㄍㄈ)", "林婉茹(ㄌㄨㄖ)", "林智偉(ㄌㄓㄨ)", "林華盛(ㄌㄏㄕ)", "林瑞華(ㄌㄖㄏ)", "林瑞燿(ㄌㄖㄧ)", "林嘉信(ㄌㄐㄒ)", 
    "邱保銘(ㄑㄅㄇ)", "邱信培(ㄑㄒㄆ)", "徐琮凱(ㄒㄘㄎ)", "柯富強(ㄎㄈㄑ)", "洪偉倫(ㄏㄨㄌ)", "胡中興(ㄏㄓㄒ)", "胡耀仁(ㄏㄧㄖ)", "夏進通(ㄒㄐㄊ)", "涂旻聖(ㄊㄇㄕ)", "張進源(ㄓㄐㄩ)", 
    "張文男(ㄓㄨㄋ)", "張世明(ㄓㄕㄇ)", "張仕欣(ㄓㄕㄒ)", "張正中(ㄓㄓㄓ)", "張百江(ㄓㄅㄐ)", "張孟哲(ㄓㄇㄓ)", "張錦升(ㄓㄐㄕ)", "張勝富(ㄓㄕㄈ)", "張路雄(ㄓㄌㄒ)", "張鈞銘(ㄓㄐㄇ)", 
    "張仕宗(ㄓㄕㄗ)", "張聰捷(ㄓㄘㄐ)", "梁善鈞(ㄌㄕㄐ)", "郭政富(ㄍㄓㄈ)", "陳文政(ㄔㄨㄓ)", "陳仕明(ㄔㄕㄇ)", "陳怡如(ㄔㄧㄖ)", "陳瑜玲(ㄔㄩㄌ)", "陳冠良(ㄔㄍㄌ)", "陳春文(ㄔㄔㄨ)", 
    "陳敏訓(ㄔㄇㄒ)", "陳盛宏(ㄔㄕㄏ)", "陳進忠(ㄔㄐㄓ)", "陳佳忠(ㄔㄐㄓ)", "曾雅惠(ㄗㄧㄏ)", "游正豪(ㄧㄓㄏ)", "游育民(ㄧㄩㄇ)", "游振和(ㄧㄓㄏ)", "黃世政(ㄏㄕㄓ)", "黃明興(ㄏㄇㄒ)", 
    "黃信隆(ㄏㄒㄌ)", "黃冠霖(ㄏㄍㄌ)", "黃堃珉(ㄏㄎㄇ)", "黃賀進(ㄏㄏㄐ)", "黃麗萍(ㄏㄌㄆ)", "楊俊逸(ㄧㄐㄧ)", "魏捷祥(ㄨㄐㄒ)", "楊閔森(ㄧㄇㄙ)", "廖崇凱(ㄌㄔㄎ)", "廖竣傑(ㄌㄐㄐ)", 
    "廖致維(ㄌㄓㄨ)", "熊致堯(ㄒㄓㄧ)", "蔡芷纭(ㄘㄓㄩ)", "蔡玉雯(ㄘㄩㄨ)", "蔡榮峰(ㄘㄖㄈ)", "鄭銘宗(ㄓㄇㄗ)", "蕭大勛(ㄒㄉㄒ)", "蕭鈺潔(ㄒㄩㄐ)", "謝仁政(ㄒㄖㄓ)", "魏志龍(ㄨㄓㄌ)", 
    "藍慧真(ㄌㄏㄓ)", "詹獻睿(ㄓㄒㄖ)", "詹孟杰(ㄓㄇㄐ)", "曾舜麟(ㄗㄕㄌ)", "謝明佑(ㄒㄇㄧ)", "羅靜宜(ㄌㄐㄧ)", "姜小平(ㄐㄒㄆ)", "莊麗雪(ㄓㄌㄒ)", "吳慧芳(ㄨㄏㄈ)", "蔡瑞賓(ㄘㄖㄅ)", 
    "張上觀(ㄓㄕㄍ)", "陳素貞(ㄔㄙㄓ)", "吳建興(ㄨㄐㄒ)", "呂昇印(ㄌㄕㄧ)", "孟繁光(ㄇㄈㄍ)", "林坤茂(ㄌㄎㄇ)", "邱榮家(ㄑㄖㄐ)", "徐偉欽(ㄒㄨㄑ)", "鄭淵太(ㄓㄩㄊ)", "張志州(ㄓㄓㄓ)", 
    "張富山(ㄓㄈㄕ)", "陳珀升(ㄔㄆㄕ)", "陳溪宗(ㄔㄒㄗ)", "陳瑋楊(ㄔㄨㄧ)", "曾明雄(ㄗㄇㄒ)", "詹憲國(ㄓㄒㄍ)", "廖宏輝(ㄌㄏㄏ)", "廖翊均(ㄌㄧㄐ)", "劉邦杰(ㄌㄅㄐ)", "劉明煌(ㄌㄇㄏ)", 
    "蔡榮祥(ㄘㄖㄒ)", "蔡榮華(ㄘㄖㄏ)", "周智勤(ㄓㄓㄑ)", "謝志忠(ㄒㄓㄓ)", "涂欽耀(ㄊㄑㄧ)", "張志仲(ㄓㄓㄓ)", "蕭森巍(ㄒㄙㄨ)", "張銀恭(ㄓㄧㄍ)", "王正錄(ㄨㄓㄌ)", "曾建勳(ㄗㄐㄒ)", 
    "黃智煒(ㄏㄓㄨ)", "劉海森(ㄌㄏㄙ)", "賴南君(ㄌㄋㄐ)", "劉權漢(ㄌㄑㄏ)", "游伊君(ㄧㄧㄐ)", "陳勇志(ㄔㄩㄓ)", "陳文田(ㄔㄨㄊ)", "林秋雄(ㄌㄑㄒ)", "賴永昌(ㄌㄩㄔ)", "劉煉騰(ㄌㄌㄊ)", 
    "林錦志(ㄌㄐㄓ)", "林明忠(ㄌㄇㄓ)", "張哲誠(ㄓㄓㄔ)", "詹昆學(ㄓㄎㄒ)", "陳彥宏(ㄔㄧㄏ)", "許馨云(ㄒㄒㄩ)", "張橋語(ㄓㄑㄩ)"
]

# --- 多重排序演算法 ---
def get_sort_key(raw_str):
    """
    第一順位：字首注音
    第二順位：字首國字 (讓同字群聚)
    第三順位：第二字注音
    第四順位：第二字國字...以此類推
    """
    name, z = raw_str.split('(')
    z = z.rstrip(')')
    key = []
    # 逐字比對注音與國字
    for i in range(max(len(name), len(z))):
        zy = z[i] if i < len(z) else ""
        char = name[i] if i < len(name) else ""
        key.extend([zy, char])
    return tuple(key)

# 執行排序並格式化
RAW_MEMBERS_SORTED = sorted(RAW_MEMBERS, key=get_sort_key)
# 下拉選單顯示格式：姓名 (注音) -> 確保打字能被搜尋到
ALL_MEMBERS_FORMATTED = [f"{m.split('(')[0]} ({m.split('(')[1]}" for m in RAW_MEMBERS_SORTED]
CLEAN_ALL_MEMBERS = [m.split('(')[0] for m in RAW_MEMBERS_SORTED]

def generate_word_report(date_str, location_str, clean_attendees):
    doc = Document()
    doc.add_heading('木工機械單車協會 - 團騎點名紀錄', 0)
    doc.add_heading(f'日期：{date_str}', level=1)
    
    if location_str:
        doc.add_heading(f'地點：{location_str}', level=2)
        
    doc.add_paragraph(f'本次參與總人數：{len(clean_attendees)} 人')
    
    attended_consultants = [name for name in clean_attendees if MEMBER_ROLES.get(name) == '顧問']
    attended_members = [name for name in clean_attendees if MEMBER_ROLES.get(name) == '隊員']
    
    doc.add_heading('出席顧問：', level=2)
    doc.add_paragraph("、".join(attended_consultants) if attended_consultants else "無")
    
    doc.add_heading('出席隊員：', level=2)
    doc.add_paragraph("、".join(attended_members) if attended_members else "無")
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def auto_adjust_excel_columns(df):
    """產出 Excel、自動調整欄寬、開啟篩選器並凍結窗格"""
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='點名總表')
        worksheet = writer.sheets['點名總表']
        
        # 1. 自動調整欄寬
        for col in worksheet.columns:
            max_length = 0
            column_letter = col[0].column_letter
            for cell in col:
                try:
                    cell_len = len(str(cell.value).encode('utf-8'))
                    if cell_len > max_length:
                        max_length = cell_len
                except:
                    pass
            worksheet.column_dimensions[column_letter].width = (max_length + 2) * 1.1
            
        # 2. 開啟全表篩選器 (包含「身份」欄位)
        worksheet.auto_filter.ref = worksheet.dimensions
        
        # 3. 凍結首列與前三欄 (編號、身份、姓名)
        worksheet.freeze_panes = 'D2'
            
    return excel_buffer.getvalue()

# 選定人員即簽到的連動函數
def on_person_select():
    selected = st.session_state.person_selector
    if selected != "--- 請點選或輸入注音搜尋 ---":
        # 簽到時，只擷取乾淨的姓名 (捨去括號與注音)
        clean_name = selected.split(' (')[0]
        if clean_name not in st.session_state.attendees:
            st.session_state.attendees.append(clean_name)
            st.session_state.report_generated = False 
    st.session_state.person_selector = "--- 請點選或輸入注音搜尋 ---"

st.title("🚴‍♂️ 車隊團騎點名系統")

if not st.session_state.setup_complete:
    st.info("💡 請先完成下方設定。確認後介面會自動鎖定並進入點名模式。")
    
    st.markdown("### 步驟一：填寫活動資訊")
    selected_date = st.date_input("📅 點擊開啟月曆選擇日期：", datetime.today())
    event_location = st.text_input("📍 輸入本次團騎地點 (例如：鳳凰山、136縣道...)")
    
    st.markdown("### 步驟二：載入資料")
    mode = st.radio("請選擇資料載入方式：", ["上傳舊有總表接續點名", "建立全新總表"])
    
    temp_df = None
    if mode == "上傳舊有總表接續點名":
        uploaded_file = st.file_uploader("📂 選擇先前的 Excel 總表", type=["xlsx"])
        if uploaded_file is not None:
            temp_df = pd.read_excel(uploaded_file)
            st.success("✅ 舊表單載入成功！")
    else:
        roles = [MEMBER_ROLES.get(m, "未知") for m in CLEAN_ALL_MEMBERS]
        temp_df = pd.DataFrame({"身份": roles, "姓名": CLEAN_ALL_MEMBERS})
        st.info("🆕 將建立全新表單。")

    if st.button("🔒 確認設定並開始點名"):
        if mode == "上傳舊有總表接續點名" and temp_df is None:
            st.warning("⚠️ 請先上傳檔案再繼續！")
        else:
            st.session_state.df = temp_df
            st.session_state.date_str = selected_date.strftime("%Y-%m-%d")
            st.session_state.location = event_location.strip()
            st.session_state.setup_complete = True
            st.rerun() 

else:
    display_event = st.session_state.date_str
    if st.session_state.location:
        display_event += f" ({st.session_state.location})"
    st.success(f"📌 目前鎖定點名活動：**{display_event}**")
    
    col1, col2 = st.columns([1.5, 1])
    
    with col1:
        st.markdown("### 🔍 步驟三：快速簽到")
        st.write("💡 打注音首字搜尋，選定後即完成簽到並從選單隱藏。")
        
        # 動態過濾：只顯示還沒簽到的人員 (比對乾淨姓名)
        AVAILABLE_OPTIONS = ["--- 請點選或輸入注音搜尋 ---"] + [
            m for m in ALL_MEMBERS_FORMATTED if m.split(' (')[0] not in st.session_state.attendees
        ]
        
        st.selectbox(
            "輸入注音或姓名關鍵字：", 
            AVAILABLE_OPTIONS, 
            key="person_selector", 
            on_change=on_person_select
        )
        
        st.markdown("---")
        st.markdown("#### 📝 修改/移除")
        # 多選框現在只顯示乾淨的姓名，完全沒有注音
        updated_attendees = st.multiselect(
            "目前已簽到清單 (點擊 x 可移除誤點人員)：", 
            st.session_state.attendees, 
            default=st.session_state.attendees,
            label_visibility="collapsed"
        )
        if updated_attendees != st.session_state.attendees:
            st.session_state.attendees = updated_attendees
            st.session_state.report_generated = False
            st.rerun()

    with col2:
        st.markdown("### 📋 即時簽到清單")
        if not st.session_state.attendees:
            st.info("尚無人員簽到")
        else:
            for i, clean_name in enumerate(st.session_state.attendees):
                role = MEMBER_ROLES.get(clean_name, "")
                st.write(f"**{i+1}.** {clean_name} ({role})")

    st.markdown("---")
    
    col_empty1, col_empty2, col_btn = st.columns([2, 1, 2])
    with col_btn:
        finish_btn = st.button("💾 點名結束！", use_container_width=True)

    if finish_btn:
        df = st.session_state.df.copy() 
        date_str = st.session_state.date_str
        location_str = st.session_state.location
        
        event_col_name = f"{date_str} {location_str}".strip()
        
        if not st.session_state.attendees:
            st.error("⚠️ 目前沒有任何人簽到！")
        elif event_col_name in df.columns and not df[df[event_col_name] == 'V'].empty:
            st.error(f"⚠️ {event_col_name} 已經有點名紀錄囉！")
        else:
            clean_attendees = st.session_state.attendees
            
            missing_members = [m for m in CLEAN_ALL_MEMBERS if m not in df['姓名'].values]
            if missing_members:
                new_roles = [MEMBER_ROLES.get(m, "未知") for m in missing_members]
                new_rows = pd.DataFrame({"身份": new_roles, "姓名": missing_members})
                df = pd.concat([df, new_rows], ignore_index=True)

            if '身份' not in df.columns:
                df.insert(1, '身份', df['姓名'].map(MEMBER_ROLES).fillna('未知'))

            if event_col_name not in df.columns:
                df[event_col_name] = ""
            for member in clean_attendees:
                df.loc[df['姓名'] == member, event_col_name] = "V"

            if '編號' in df.columns:
                df = df.drop(columns=['編號'])
            if '總次數' in df.columns:
                df = df.drop(columns=['總次數'])

            date_cols = [col for col in df.columns if col not in ['姓名', '身份', '編號', '總次數']]
            df['總次數'] = df[date_cols].apply(lambda x: (x == 'V').sum(), axis=1)
            
            df.insert(0, '編號', range(1, len(df) + 1))
            final_cols = ['編號', '身份', '姓名'] + date_cols + ['總次數']
            df = df[final_cols]
            
            st.session_state.final_df = df
            
            st.session_state.excel_data = auto_adjust_excel_columns(df)
            st.session_state.word_data = generate_word_report(date_str, location_str, clean_attendees)
            st.session_state.report_generated = True

    if st.session_state.report_generated:
        st.success(f"🎉 已成功記錄 {len(st.session_state.attendees)} 人！您可以隨時點擊下方按鈕下載檔案。")

        file_name_suffix = st.session_state.date_str
        if st.session_state.location:
             file_name_suffix += f"_{st.session_state.location}"
             
        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button("📥 下載 Excel 總表", data=st.session_state.excel_data, file_name=f"車隊點名總表_{file_name_suffix}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        with dl_col2:
            st.download_button("📥 下載 Word 紀錄檔", data=st.session_state.word_data, file_name=f"團騎點名紀錄_{file_name_suffix}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        st.write("📊 總表預覽：")
        st.dataframe(st.session_state.final_df)

    st.markdown("---")
    if st.button("⚙️ 重新設定日期或表單"):
        st.session_state.setup_complete = False
        st.rerun()
